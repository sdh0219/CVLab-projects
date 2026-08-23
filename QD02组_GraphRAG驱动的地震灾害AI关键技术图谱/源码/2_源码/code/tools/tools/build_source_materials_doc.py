from __future__ import annotations

import csv
import json
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
MANIFEST_PATH = PROJECT_ROOT / "outputs" / "graphrag_index" / "exports" / "corpus_manifest.csv"
SUMMARY_PATH = PROJECT_ROOT / "outputs" / "graphrag_index" / "index_summary.json"
SUBMISSION_MANIFEST_PATH = PROJECT_ROOT / "outputs" / "graphrag_index" / "exports" / "submission_manifest.json"
OUTPUT_PATH = PROJECT_ROOT / "docs" / "资料来源与用途说明.docx"


TYPE_LABELS = {
    "case": "灾害案例",
    "paper": "论文",
    "patent": "专利",
    "policy": "政策",
    "project": "项目",
    "report": "报告",
    "standard": "标准",
}

TYPE_ORDER = ["case", "paper", "patent", "policy", "project", "report", "standard"]

TYPE_USES = {
    "case": "用于形成 Case 节点与 VALIDATED_IN（案例验证）关系，支撑真实地震情境下的技术验证。",
    "paper": "用于抽取 AI 技术、模型、任务和数据依赖，支撑 USES_MODEL（采用模型）、DEPENDS_ON（依赖数据）和 SOLVES（解决任务）关系。",
    "patent": "用于补充工程方法、预警装置、恢复预测等技术路线，作为关键技术方案和应用边界证据。",
    "policy": "用于形成 Policy 节点与 REQUIRED_BY（政策要求）关系，支撑应急管理、风险治理和合规约束说明。",
    "project": "用于提供工程系统、平台或开源工具案例，支撑技术落地路径、数据条件和应用场景描述。",
    "report": "用于提供宏观评估、风险治理、基础设施韧性和 AI 应用背景，支撑社区报告和关键技术评估。",
    "standard": "用于提供数据接口、公共预警、风险管理和应急流程规范，支撑标准化约束与流程建模。",
}

SOURCE_TYPE_SUMMARY = {
    "case": "ReliefWeb、World Bank、GFDRR",
    "paper": "Remote Sensing、arXiv、UC Berkeley、GEM Foundation、GFDRR",
    "patent": "Google Patents",
    "policy": "UNDRR、中国国务院、FEMA、NIST、日本内阁府",
    "project": "USGS、UC Berkeley、NASA、GEM Foundation、FEMA、HOT OSM、xView2",
    "report": "GFDRR、World Bank、JRC、UNDRR、FEMA、CDRI",
    "standard": "ISO、OGC、OASIS",
}

EXPORT_DESCRIPTIONS = {
    "graph_nodes.csv": "图谱节点表；用于检查实体类型、社区编号和节点属性。",
    "graph_edges_neo4j.csv": "Neo4j 导入边表；用于把关系结构导入图数据库。",
    "evidence_edges.csv": "证据边表；用于追溯每条关系对应的来源文档、证据片段和置信度。",
    "community_reports.csv": "社区报告表；用于说明关键技术社区及其证据基础。",
    "atlas.graphml": "GraphML 图谱文件；用于 Gephi、Cytoscape 等外部工具分析。",
    "neo4j_import.cypher": "Neo4j 导入脚本；用于批量创建节点和关系。",
    "atlas_quality_report.json": "质量报告 JSON；用于校验文档、实体、关系、社区和待复核项。",
    "expert_review_log.csv": "专家复核记录表；用于后续人工审核和关系确认。",
    "corpus_manifest.csv": "语料来源清单；用于记录 49 条资料的来源、URL、年份、区域和质量状态。",
    "technology_nodes.csv": "技术实体节点表；用于抽取和展示 AI 关键技术清单。",
    "key_technology_assessment.csv": "关键技术评分表；用于报告中说明技术成熟度和证据强度。",
    "qa_evaluation_set.csv": "问答评测种子集；用于检验 GraphRAG 检索与回答质量。",
    "topology_communities.csv": "拓扑社区分析表；用于解释节点社群结构。",
    "uncertainty_report.md": "冲突与不确定性报告；用于提示证据不足和需复核内容。",
}

INTERNAL_ASSETS = [
    ("data/corpus/sample/*.md", "地震专题中文整理稿", "作为 GraphRAG 索引输入，提供实体、关系、声明和证据片段抽取材料。"),
    ("data/corpus/README.md", "语料来源说明", "说明 49 条资料的类型分布、来源边界和正式引用注意事项。"),
    ("config/ontology.json", "领域本体配置", "定义实体类型、关系类型、抽取约束和地震灾害技术图谱建模口径。"),
    ("graphrag_atlas/", "GraphRAG 索引代码", "完成文档切分、规则抽取、关系生成、社区报告和检索问答。"),
    ("tools/build_earthquake_corpus.py", "语料构建脚本", "生成或刷新地震专题语料。"),
    ("tools/enhance_atlas_outputs.py", "成果增强脚本", "生成 Neo4j、GraphML、质量报告、评估集和复核表。"),
    ("tools/publish_atlas_frontend.py", "前端发布脚本", "把索引成果发布为前端可读取的 JSON 快照。"),
    ("public/atlas/atlas_frontend.json", "前端图谱快照", "支撑交互图谱、节点展开、边聚合、详情面板和证据片段展示。"),
    ("docs/", "项目说明文档", "记录方案设计、运行方式、文件结构、前端验证和阶段性工作。"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_fixed(table, widths_cm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        set_row_cant_split(row)
        for index, width in enumerate(widths_cm):
            if index >= len(row.cells):
                continue
            cell = row.cells[index]
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_mixed_paragraph(doc: Document, label: str, text: str, style: str = "Body Text") -> None:
    paragraph = doc.add_paragraph(style=style)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True, color="1F4D78")
    text_run = paragraph.add_run(text)
    set_run_font(text_run)


def add_table_header(table, headers: list[str]) -> None:
    set_row_repeat_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=9, bold=True, color="0B2545")


def style_table_text(table, body_size: float = 8.5) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9 if row_index == 0 else body_size, bold=run.bold)


def read_manifest() -> list[dict[str, str]]:
    with MANIFEST_PATH.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def source_use(row: dict[str, str]) -> str:
    base = TYPE_USES.get(row["source_type"], "用于提供项目语料和证据追溯。")
    return f"{base}本条资料主要支撑“{row.get('title', '').strip()}”对应的证据链。"


def add_source_type_section(doc: Document, source_type: str, rows: list[dict[str, str]]) -> None:
    label = TYPE_LABELS[source_type]
    doc.add_heading(f"{label}资料", level=2)
    add_mixed_paragraph(
        doc,
        "用途说明：",
        TYPE_USES[source_type],
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_fixed(table, [3.5, 4.1, 3.0, 5.4])
    add_table_header(table, ["资料", "来源与链接", "技术主线", "在项目中的用处"])
    for row in rows:
        row_obj = table.add_row()
        set_row_cant_split(row_obj)
        cells = row_obj.cells
        cells[0].text = f"{row['doc_id']}\n{row['title']}\n{TYPE_LABELS[row['source_type']]}，{row['year']}，{row['region']}"
        cells[1].text = f"{row['source_name']}\n{row['source_url']}\n本地整理稿：{row['path']}"
        cells[2].text = infer_technology_line(row["doc_id"], row["path"])
        cells[3].text = source_use(row)
    style_table_text(table, body_size=7.5)


def infer_technology_line(doc_id: str, path_text: str) -> str:
    path = PROJECT_ROOT / path_text
    if not path.exists():
        return "地震灾害 AI 技术图谱证据"
    text = path.read_text(encoding="utf-8")
    marker = "主要支撑 "
    if marker in text:
        after = text.split(marker, 1)[1]
        return after.split(" 技术社区", 1)[0].strip()
    fallback_map = {
        "warning": "地震早期预警",
        "damage": "震后建筑损毁识别",
        "lifeline": "生命线震损风险传播GNN",
        "risk": "地震风险时空预测",
        "qa": "地震应急辅助决策",
        "hazus": "震后灾情快速评估",
    }
    for key, value in fallback_map.items():
        if key in doc_id:
            return value
    return "地震灾害 AI 技术图谱证据"


def build_doc() -> None:
    rows = read_manifest()
    summary = load_json(SUMMARY_PATH)
    submission_manifest = load_json(SUBMISSION_MANIFEST_PATH)
    type_counts = Counter(row["source_type"] for row in rows)
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        grouped[row["source_type"]].append(row)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Body Text"].font.name = "Calibri"
    styles["Body Text"].font.size = Pt(10.5)
    styles["Body Text"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI防灾减灾关键技术图谱构建资料来源与用途说明")
    set_run_font(title_run, size=18, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("以地震灾害为例的 GraphRAG 项目资料清单")
    set_run_font(sub_run, size=11, color="555555")

    add_mixed_paragraph(doc, "项目目录：", str(PROJECT_ROOT))
    add_mixed_paragraph(doc, "生成日期：", date.today().isoformat())
    add_mixed_paragraph(doc, "资料范围：", f"{summary.get('documents', len(rows))} 条地震专题中文整理稿，覆盖 7 类来源；每条均保留 source_url、年份、区域和本地整理稿路径。")
    add_mixed_paragraph(doc, "使用边界：", "本项目语料是面向实体、关系和证据抽取的中文整理稿，不是原始论文、网页或标准全文；正式引用时应回到 source_url 核对原始来源。")

    doc.add_heading("一、资料体系总览", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_fixed(table, [2.2, 1.5, 4.1, 8.2])
    add_table_header(table, ["资料类型", "数量", "主要来源", "项目用途"])
    for source_type in TYPE_ORDER:
        row_obj = table.add_row()
        set_row_cant_split(row_obj)
        cells = row_obj.cells
        cells[0].text = TYPE_LABELS[source_type]
        cells[1].text = str(type_counts[source_type])
        cells[2].text = SOURCE_TYPE_SUMMARY[source_type]
        cells[3].text = TYPE_USES[source_type]
    style_table_text(table, body_size=8.5)

    doc.add_heading("二、索引与成果使用范围", level=1)
    add_mixed_paragraph(doc, "索引规模：", f"文档 {summary.get('documents')}，文本块 {summary.get('chunks')}，实体 {summary.get('entities')}，声明 {summary.get('claims')}，关系 {summary.get('relations')}，技术社区 {summary.get('communities')}。")
    add_mixed_paragraph(doc, "抽取方式：", f"当前索引摘要记录的抽取器为 {summary.get('extractor')}；项目同时保留 LLM 和 hybrid 抽取入口，但本次已落盘成果以当前索引摘要为准。")
    add_mixed_paragraph(doc, "资料流向：", "语料 Markdown 经 GraphRAG 管线切分为 documents/chunks，再抽取 entities、claims、relations，进一步生成社区报告、证据边、Neo4j/GraphML 导出和前端 atlas 快照。")

    doc.add_heading("三、项目内部资料与用途", level=1)
    internal_table = doc.add_table(rows=1, cols=3)
    internal_table.style = "Table Grid"
    set_table_fixed(internal_table, [4.1, 3.2, 8.7])
    add_table_header(internal_table, ["本地资料或目录", "资料性质", "用途"])
    for path, nature, use in INTERNAL_ASSETS:
        row_obj = internal_table.add_row()
        set_row_cant_split(row_obj)
        cells = row_obj.cells
        cells[0].text = path
        cells[1].text = nature
        cells[2].text = use
    style_table_text(internal_table, body_size=8.5)

    doc.add_heading("四、主要导出成果文件与用途", level=1)
    export_files = submission_manifest.get("export_files", [])
    core_names = [
        "graph_nodes.csv",
        "graph_edges_neo4j.csv",
        "evidence_edges.csv",
        "community_reports.csv",
        "atlas.graphml",
        "neo4j_import.cypher",
        "atlas_quality_report.json",
        "expert_review_log.csv",
        "corpus_manifest.csv",
        "technology_nodes.csv",
        "key_technology_assessment.csv",
        "qa_evaluation_set.csv",
        "topology_communities.csv",
        "uncertainty_report.md",
    ]
    export_by_name = {item["name"]: item for item in export_files}
    export_table = doc.add_table(rows=1, cols=3)
    export_table.style = "Table Grid"
    set_table_fixed(export_table, [4.1, 2.0, 9.9])
    add_table_header(export_table, ["文件", "大小", "用途"])
    for name in core_names:
        if name not in export_by_name:
            continue
        row_obj = export_table.add_row()
        set_row_cant_split(row_obj)
        cells = row_obj.cells
        cells[0].text = f"outputs/graphrag_index/exports/{name}"
        cells[1].text = f"{int(export_by_name[name]['bytes']) / 1024:.1f} KB"
        cells[2].text = EXPORT_DESCRIPTIONS.get(name, "项目导出成果文件，用于图谱分析、校验或提交。")
    style_table_text(export_table, body_size=8.5)

    doc.add_heading("五、外部资料来源明细", level=1)
    add_mixed_paragraph(doc, "说明：", "以下 49 条记录来自 corpus_manifest.csv，并与 data/corpus/sample 下的 Markdown 整理稿对应。表中“在项目中的用处”指该资料进入图谱后的主要建模作用。")
    for source_type in TYPE_ORDER:
        add_source_type_section(doc, source_type, grouped[source_type])

    doc.add_heading("六、引用与提交建议", level=1)
    add_mixed_paragraph(doc, "正式引用：", "报告正文引用外部资料时，优先使用表中 source_url 对应的原始网页、论文页、专利页或标准页，不把本地中文整理稿当作原始出处。")
    add_mixed_paragraph(doc, "成果提交：", "建议提交 README.md、docs 说明文档、data/corpus、config/ontology.json、graphrag_atlas、tools、outputs/graphrag_index、public/atlas 和最终提交包；不建议提交 node_modules、dist、build、.wrangler、.idea 和日志文件。")
    add_mixed_paragraph(doc, "复核要求：", "expert_review_log.csv 和 expert_review_priority.csv 为待复核结构，当前 pending 状态表示后续仍需人工或专家确认。")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("AI防灾减灾关键技术图谱构建资料来源与用途说明")
    set_run_font(footer_run, size=9, color="555555")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
