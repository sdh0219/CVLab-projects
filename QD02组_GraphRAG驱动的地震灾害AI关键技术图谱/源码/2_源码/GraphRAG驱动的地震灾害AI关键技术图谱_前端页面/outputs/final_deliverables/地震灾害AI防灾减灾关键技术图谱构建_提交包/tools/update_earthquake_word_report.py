from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def main() -> None:
    parser = argparse.ArgumentParser(description="Update the external DOCX report to the earthquake-focused scope.")
    default_path = (
        Path(__file__).resolve().parents[1].parent
        / "3.结果"
        / "04_报告"
        / "AI防灾减灾知识图谱构建技术方案报告.docx"
    )
    parser.add_argument("--path", default=str(default_path))
    args = parser.parse_args()

    path = Path(args.path)
    doc = Document(path)
    update_paragraphs(doc)
    update_tables(doc)
    apply_chinese_font(doc)
    doc.save(path)
    print(path)


def update_paragraphs(doc: Document) -> None:
    updates = {
        0: "地震灾害AI防灾减灾关键技术图谱构建技术方案报告",
        1: "1. 项目背景与意义",
        2: (
            "地震灾害风险识别、地震早期预警、震后应急响应和震后损毁评估涉及地震台网、"
            "强震动记录、遥感影像、建筑物足迹、生命线设施和灾情文本等多源信息。AI 技术"
            "能够提升地震信息处理效率，但不同技术的输入数据、适用阶段、输出结果和落地条件"
            "差异明显。构建地震灾害AI关键技术图谱，有助于把分散资料组织为可追溯、可比较、"
            "可展示的证据网络。"
        ),
        3: (
            "本项目围绕地震早期预警、震后建筑损毁识别、遥感震损智能解译、地震多模态灾情"
            "理解、地震应急辅助决策、震后救援调度优化、生命线震损风险传播和地震证据图谱 "
            "GraphRAG 等技术，整理地震案例、抽取节点关系、绘制可视化成果，并形成数据包、"
            "源码、图件、技术方案报告和演示材料。"
        ),
        4: "2. 问题定义与项目目标",
        5: (
            "本项目不开展新的深度模型训练实验，项目核心是地震专题资料整理、案例分析、节点"
            "关系建模和知识图谱可视化。目标是形成一套能够说明地震灾害AI关键技术、核心能力、"
            "应用阶段、输入数据、技术方法、输出结果、代表案例、适用边界和参考依据的结构化成果。"
        ),
        6: (
            "最终成果由地震专题语料、源码、图件、技术方案报告和演示材料共同构成。知识图谱由"
            "节点表、关系表、案例表、参考资料表和可视化成果图共同表达，单张图片不能代表完整成果。"
        ),
        7: "3. 数据来源与资料整理",
        8: (
            "资料来源包括课程项目提交要求、资料准备清单、AI 防灾减灾资料收集与分析文档、技术框架图、"
            "技术图谱参考图、UNDRR 仙台框架、ShakeAlert、MyShake、OpenQuake、Hazus、xBD 数据集、"
            "NASA Disasters、GFDRR 震后评估资料、ISO/OASIS 预警标准和地震应急管理公开案例。"
        ),
        9: (
            "资料整理过程将文本、图像和表格材料统一转化为地震技术案例字段，包括技术类型、核心能力、"
            "地震阶段、代表案例、输入数据、技术方法、输出结果、解决问题、优势、难点和参考依据。"
        ),
        10: "4. 地震灾害AI关键技术图谱定义",
        11: (
            "本项目中的地震灾害AI关键技术图谱，是面向震前监测感知、地震风险评估、地震早期预警、"
            "震后应急调度和震后损毁评估等任务，围绕机器学习、计算机视觉、自然语言处理、强化学习、"
            "图神经网络和大模型等 AI 技术构建的“技术类型-核心能力-地震阶段-输入数据-技术方法-"
            "输出结果-代表案例-适用边界-参考依据”多层级知识网络。"
        ),
        12: "5. 地震灾害AI技术分类体系",
        13: (
            "机器学习主要面向地震台网数据驱动的震相识别、地震早期预警和震害风险预测。计算机视觉"
            "主要面向震后建筑损毁识别、遥感震损智能解译和无人机震损核查。自然语言处理主要面向"
            "地震灾情信息抽取、应急文本分类、灾情舆情风险识别和灾情摘要生成。强化学习与数字孪生"
            "主要面向震后救援路径规划、物资调度和生命线设施恢复推演。大模型主要面向地震应急知识助手、"
            "预案与救援方案生成、多模态灾情辅助研判和证据链问答。"
        ),
        14: "6. 技术路线与图谱构建方法",
        15: (
            "技术路线包括地震资料收集、案例整理、字段标准化、节点抽取、关系定义、图谱绘制、结果分析"
            "和成果输出。节点建模以中心主题、技术类型、核心能力、地震阶段、代表案例、输入数据、技术方法、"
            "输出结果、解决问题、优势、难点、发展趋势、质量控制和参考依据为主要节点类型。关系建模覆盖"
            "包含技术、具备能力、面向地震、应用于、依赖数据、采用方法、生成结果、解决问题、具有优势、"
            "面临难点、来源支持、发展方向和质量控制。"
        ),
        17: "图 1 知识图谱构建流程图",
        18: "7. 地震技术案例总表分析",
        19: (
            "地震技术案例总表共整理 16 条案例，覆盖地震早期预警、震后损毁评估、地震灾情文本处理、"
            "震后应急调度和大模型辅助决策等方向。案例总表支撑统计图、矩阵图、节点表和关系表生成。"
        ),
        21: "图 2 技术案例数量统计图",
        23: "8. 地震图谱节点与关系建模",
        24: (
            "当前 GraphRAG 地震专题索引包含 49 篇整理语料、58 个实体节点、202 条声明、792 条证据关系"
            "和 10 个技术社区，能够表达从地震技术类型到数据依赖、模型方法、地震案例、政策标准和限制条件"
            "的完整链路。"
        ),
        26: "9. 自然语言处理模块重点分析",
        27: (
            "自然语言处理模块是课程特色重点。该模块处理社交媒体、新闻、求助帖、应急通报和现场记录等"
            "非结构化文本，将其中的地点、时间、地震事件、救援需求、紧急程度、舆情热点和态势摘要转化为"
            "结构化结果。"
        ),
        29: "10. 地震技术图谱结果展示",
        30: (
            "地震灾害AI关键技术图谱可视化成果图以技术类型为行，以核心能力、地震阶段、数据与方法、结果与"
            "边界、参考依据为列，展示从技术实体到应用证据的完整链路。图件底部补充知识链路，使图面逻辑与"
            "节点表、关系表保持一致。"
        ),
        32: "图 3 地震灾害AI关键技术图谱可视化成果图",
        34: "图 4 地震技术类型-能力-场景-案例矩阵图",
        35: "11. 地震场景下不同技术适用问题分析",
        36: (
            "不同 AI 技术适用于不同地震信息形态。机器学习适合地震台网和历史事件数据较充足的预测评估问题；"
            "计算机视觉适合遥感影像、视频和无人机影像支撑的震损识别问题；自然语言处理适合灾情文本和公众"
            "表达处理；强化学习和数字孪生适合震后救援推演和方案比选；大模型适合知识检索、材料生成和多源"
            "信息辅助研判。"
        ),
        38: "图 5 适用问题与落地难点图",
        39: "12. 应用价值",
        40: (
            "本项目的应用价值体现在三方面。第一，形成可追溯的地震专题数据基础，使案例、节点、关系和参考"
            "依据之间可以相互核对。第二，形成可复现的源码流程，使统计图、流程图、图谱图、矩阵图和难点图"
            "能够由数据表自动生成。第三，形成可展示的课程成果包，使报告和讲稿能够围绕同一套地震专题数据"
            "结构展开。"
        ),
        41: "13. 存在问题与改进方向",
        42: (
            "项目仍存在资料来源粒度不完全一致、不同技术案例成熟度差异较大、部分地震场景缺少统一评价指标等"
            "问题。后续可从三个方向改进：一是接入更多标准化地震台网、遥感震损和震后评估数据集；二是增加"
            "实体消歧、关系权重和证据评分；三是建设交互式知识图谱页面，支持按技术类型、地震阶段和参考依据"
            "筛选。"
        ),
        43: "14. 总结",
        44: (
            "本报告完成了地震灾害AI关键技术图谱的定义、资料整理、技术分类、案例分析、节点关系建模和可视化"
            "展示。成果包能够支撑课程汇报、技术选型说明和后续扩展。自然语言处理模块作为课程重点，展示了"
            "地震灾情文本从抽取、分类、识别到摘要生成的完整应用链路。"
        ),
        45: "15. 参考资料",
    }
    for index, text in updates.items():
        if index < len(doc.paragraphs):
            doc.paragraphs[index].text = text


def update_tables(doc: Document) -> None:
    case_rows = [
        ["案例编号", "技术类型", "地震场景", "输出结果", "参考依据"],
        ["EQW-01", "机器学习", "地震台网秒级预警", "预警时间窗、震中位置、预估烈度", "R04、R05"],
        ["EQW-02", "机器学习", "手机众包地震预警", "震动触发事件、区域预警、置信度", "R05、R06"],
        ["EQW-03", "时空预测", "地震风险评估", "风险等级、重点防范区域、损失预估", "R07、R08"],
        ["CV-01", "计算机视觉", "震后建筑损毁识别", "建筑损毁等级、空间分布、核查清单", "R09、R10"],
        ["CV-02", "遥感智能解译", "遥感震损快速评估", "震损斑块、受影响建筑、评估图层", "R10、R11"],
        ["CV-03", "无人机视觉", "现场震损核查", "重点点位、损毁证据、巡查路径", "R11、R12"],
        ["NLP-01", "自然语言处理", "地震灾情信息抽取", "地点、时间、地震事件、救援需求", "R13、R14"],
        ["NLP-02", "自然语言处理", "地震应急文本分类", "紧急程度、业务类别、分发标签", "R13、R15"],
        ["NLP-03", "自然语言处理", "震后舆情风险识别", "舆情热点、情绪趋势、谣言风险", "R15、R16"],
        ["NLP-04", "自然语言处理", "地震灾情摘要生成", "态势简报、信息要点、处置建议", "R14、R17"],
        ["SIM-01", "数字孪生", "震后救援推演", "可达性评估、路径方案、资源瓶颈", "R18、R19"],
        ["SIM-02", "图神经网络", "生命线震损风险传播", "级联风险、关键节点、恢复优先级", "R19、R20"],
        ["SIM-03", "强化学习", "震后物资调度优化", "队伍调度、车辆路径、补给计划", "R18、R21"],
        ["LLM-01", "大模型", "地震应急知识问答", "证据链答案、参考来源、复核提示", "R22、R23"],
        ["LLM-02", "大模型", "预案与救援方案生成", "任务清单、行动方案、风险提示", "R22、R24"],
        ["LLM-03", "多模态大模型", "多源灾情辅助研判", "图文证据、态势解释、待核查项", "R23、R24"],
    ]
    node_rows = [
        ["节点或成果", "数量"],
        ["地震专题语料", "49"],
        ["实体节点", "58"],
        ["证据关系", "792"],
        ["声明证据", "202"],
        ["技术社区", "10"],
        ["拓扑社区", "3"],
        ["灾害类型", "1"],
        ["应用场景", "5"],
        ["AI技术节点", "10"],
        ["模型方法", "8"],
        ["数据类型", "7"],
        ["任务类型", "7"],
        ["地震案例", "8"],
        ["政策标准", "5"],
        ["限制条件", "7"],
    ]
    nlp_rows = [
        ["NLP 任务", "输入文本", "技术方法", "输出结果", "地震应急价值"],
        ["地震灾情信息抽取", "微博、新闻、求助帖、现场记录", "命名实体识别、BERT-BiLSTM-CRF", "地点、时间、地震事件、救援需求", "支撑震后态势汇聚"],
        ["地震应急文本分类", "通报、短信、群聊文本", "BERT、文本分类、情感分析", "紧急程度、业务类别、分发标签", "支撑信息分级和部门路由"],
        ["震后舆情风险识别", "微博、评论、转发文本", "主题模型、情感分析、谣言识别", "舆情热点、情绪趋势、风险话题", "支撑舆情研判和信息发布"],
        ["地震灾情摘要生成", "灾情文本、新闻报道、现场记录", "文本摘要、信息融合、生成式模型", "灾情摘要、态势简报、信息要点", "支撑快速形成汇报材料"],
    ]
    refs = [
        ["编号", "资料标题", "作者或机构", "年份", "来源类型", "对应场景"],
        ["R01", "AI+防灾减灾学生实践项目提交要求", "课程材料", "2026", "课程要求", "课程实践成果"],
        ["R02", "AI 防灾减灾技术图谱项目资料准备清单", "课程材料", "2026", "课程材料", "图谱构建资料体系"],
        ["R03", "AI 防灾减灾资料收集与分析", "课程材料", "2026", "课程材料", "地震专题案例整理"],
        ["R04", "ShakeAlert Earthquake Early Warning", "USGS", "2024", "示范项目", "地震早期预警"],
        ["R05", "MyShake Smartphone Earthquake Early Warning", "UC Berkeley", "2024", "示范项目", "手机众包预警"],
        ["R06", "Earthquake Early Warning Patent", "Google Patents", "2014", "专利", "地震预警发布"],
        ["R07", "OpenQuake Engine", "GEM Foundation", "2024", "开源项目", "地震风险评估"],
        ["R08", "Hazus Earthquake Loss Estimation", "FEMA", "2024", "工程工具", "震害损失估计"],
        ["R09", "BDD-Net Building Damage Detection", "Remote Sensing", "2020", "论文", "震后建筑损毁识别"],
        ["R10", "xBD Dataset for Building Damage Assessment", "arXiv", "2019", "数据集论文", "遥感震损评估"],
        ["R11", "NASA Disasters Program", "NASA", "2024", "示范项目", "震后遥感评估"],
        ["R12", "Humanitarian Mapping for Earthquake Response", "HOT", "2024", "项目", "震后现场核查"],
        ["R13", "Crisis Text Analytics for Emergency Response", "公开研究资料", "2024", "论文资料", "灾情文本抽取"],
        ["R14", "Large Language Models for Emergency Knowledge Assistance", "GFDRR", "2020", "报告", "证据链问答"],
        ["R15", "Responsible AI for Disaster Risk Management", "GFDRR", "2020", "报告", "AI治理边界"],
        ["R16", "NIST AI Risk Management Framework", "NIST", "2023", "政策框架", "模型治理"],
        ["R17", "Artificial Intelligence Approaches for Disaster Risk Management", "JRC", "2025", "报告", "AI应急应用"],
        ["R18", "ISO 22320 Incident Management", "ISO", "2018", "标准", "地震应急管理"],
        ["R19", "Lifelines: The Resilient Infrastructure Opportunity", "World Bank and GFDRR", "2019", "报告", "生命线韧性"],
        ["R20", "Global Infrastructure Resilience Report", "CDRI", "2023", "报告", "基础设施韧性"],
        ["R21", "National Incident Management System", "FEMA", "2023", "政策", "震后协调调度"],
        ["R22", "Sendai Framework for Disaster Risk Reduction", "UNDRR", "2015", "国际框架", "地震风险治理"],
        ["R23", "Common Alerting Protocol", "OASIS", "2010", "标准", "预警消息互操作"],
        ["R24", "ISO 22322 Public Warning", "ISO", "2022", "标准", "公众预警发布"],
        ["R25", "Turkiye Earthquakes Recovery and Reconstruction Assessment", "GFDRR", "2023", "地震案例", "震后恢复评估"],
        ["R26", "Japan Noto Peninsula Earthquake Case", "ReliefWeb", "2024", "地震案例", "地震预警与响应"],
        ["R27", "Nepal Earthquake Post Disaster Needs Assessment", "ReliefWeb", "2015", "地震案例", "震后需求评估"],
    ]
    fill_table(doc.tables[0], case_rows)
    fill_table(doc.tables[1], node_rows)
    fill_table(doc.tables[2], nlp_rows)
    fill_table(doc.tables[3], refs)
    for table in doc.tables:
        repeat_header(table.rows[0])
        for row in table.rows:
            prevent_row_split(row)


def fill_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    for row_index, row in enumerate(rows):
        for cell_index, value in enumerate(row):
            if cell_index < len(table.rows[row_index].cells):
                table.rows[row_index].cells[cell_index].text = value
    for row_index in range(len(rows), len(table.rows)):
        for cell in table.rows[row_index].cells:
            cell.text = ""


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def apply_chinese_font(doc: Document) -> None:
    font_name = "Microsoft YaHei"
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            if getattr(style, "font", None):
                style.font.name = font_name
                style.font.size = style.font.size or Pt(10.5)
                style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, font_name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, font_name)


def set_run_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


if __name__ == "__main__":
    main()
