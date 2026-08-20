from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "个人结课报告"
IMAGE_DIR = OUT_DIR / "generated_images"
PROJECT_TITLE = "避开危险路段的救援路径规划系统设计与实现"
REPORT_DATE = "2026年6月26日"


def load_image_font(size: int, bold: bool = False):
    font_dir = Path("C:/Windows/Fonts")
    candidates = [
        font_dir / ("msyhbd.ttc" if bold else "msyh.ttc"),
        font_dir / ("simhei.ttf" if bold else "simsun.ttc"),
        font_dir / "arial.ttf",
    ]
    for path in candidates:
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size)
            except OSError:
                pass
    return ImageFont.load_default()


def rounded_box(draw, box, fill, outline="#475569", width=2, radius=18) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_wrapped_text(draw, xy, text: str, font, fill="#0f172a", max_width=260, line_gap=6) -> None:
    x, y = xy
    for raw_line in text.splitlines() or [""]:
        line = ""
        for ch in raw_line:
            test = line + ch
            if draw.textlength(test, font=font) <= max_width:
                line = test
            else:
                draw.text((x, y), line, font=font, fill=fill)
                y += font.size + line_gap
                line = ch
        if line:
            draw.text((x, y), line, font=font, fill=fill)
            y += font.size + line_gap


def arrow(draw, start, end, fill="#334155", width=4) -> None:
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - direction * 16, y2 - 9), (x2 - direction * 16, y2 + 9)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - direction * 16), (x2 + 9, y2 - direction * 16)]
    draw.polygon(points, fill=fill)


def create_flow_image(path: Path, title: str, steps: list[str], accent="#2563eb") -> None:
    w, h = 1700, 680
    img = Image.new("RGB", (w, h), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = load_image_font(42, True)
    body_font = load_image_font(25)
    small_font = load_image_font(21)
    draw.text((70, 42), title, font=title_font, fill="#0f172a")
    draw.text((72, 98), "按照课程报告写作方式整理：先说明输入数据，再说明处理过程，最后输出可解释结果。", font=small_font, fill="#64748b")
    x0, y0 = 80, 230
    box_w, box_h, gap = 205, 155, 30
    colors = ["#dbeafe", "#dcfce7", "#fef3c7", "#fee2e2", "#ede9fe", "#e0f2fe"]
    for i, step in enumerate(steps):
        x = x0 + i * (box_w + gap)
        rounded_box(draw, (x, y0, x + box_w, y0 + box_h), colors[i % len(colors)], accent, 3)
        draw.ellipse((x + 18, y0 + 18, x + 58, y0 + 58), fill=accent)
        draw.text((x + 31, y0 + 23), str(i + 1), font=body_font, fill="white")
        draw_wrapped_text(draw, (x + 20, y0 + 76), step, body_font, "#0f172a", max_width=box_w - 40)
        if i < len(steps) - 1:
            arrow(draw, (x + box_w + 5, y0 + box_h // 2), (x + box_w + gap - 6, y0 + box_h // 2), accent, 4)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, quality=95)


def create_weight_image(path: Path) -> None:
    w, h = 1500, 760
    img = Image.new("RGB", (w, h), "#ffffff")
    draw = ImageDraw.Draw(img)
    title_font = load_image_font(42, True)
    body_font = load_image_font(27)
    small_font = load_image_font(23)
    draw.text((70, 42), "安全路径综合权重设计示意图", font=title_font, fill="#0f172a")
    boxes = [
        ("道路距离\ndistance", "#dbeafe"),
        ("灾害风险系数\nrisk_factor", "#fee2e2"),
        ("拥堵修正\ncongestion", "#fef3c7"),
        ("固定惩罚\nfixed_cost", "#dcfce7"),
    ]
    x, y = 85, 190
    for i, (label, color) in enumerate(boxes):
        bx = x + i * 335
        rounded_box(draw, (bx, y, bx + 255, y + 140), color, "#334155", 3)
        draw_wrapped_text(draw, (bx + 36, y + 36), label, body_font, max_width=190)
        if i < len(boxes) - 1:
            arrow(draw, (bx + 268, y + 70), (bx + 322, y + 70), "#64748b", 4)
    formula = "safe_weight = distance × risk_factor × (1 + congestion_weight × congestion) + fixed_cost"
    rounded_box(draw, (130, 430, 1370, 555), "#f8fafc", "#2563eb", 4)
    draw.text((175, 470), formula, font=body_font, fill="#1d4ed8")
    notes = [
        "所有权重保持非负，满足 Dijkstra 算法使用条件。",
        "风险越高，综合代价越高，算法越倾向于选择绕行道路。",
        "安全路径的最优含义是综合代价最低，不等同于距离最短。",
    ]
    for i, note in enumerate(notes):
        draw.text((170, 610 + i * 36), f"{i + 1}. {note}", font=small_font, fill="#334155")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, quality=95)


def create_result_chart(path: Path) -> None:
    data = [
        ("四川地震", 156.684, 209.05, 825, 0),
        ("北京洪水", 19.905, 22.764, 86, 1),
        ("上海火灾", 2.001, 3.725, 14, 7),
    ]
    w, h = 1500, 850
    img = Image.new("RGB", (w, h), "#ffffff")
    draw = ImageDraw.Draw(img)
    title_font = load_image_font(42, True)
    body_font = load_image_font(24)
    small_font = load_image_font(21)
    draw.text((70, 42), "三类灾害场景路径结果对比", font=title_font, fill="#0f172a")
    draw.text((72, 98), "蓝色表示普通最短路径，绿色表示安全路径；下方数字表示危险边数量。", font=small_font, fill="#64748b")
    left, top, bottom = 130, 190, 650
    max_dist = 220
    draw.line((left, top, left, bottom), fill="#94a3b8", width=2)
    draw.line((left, bottom, 1360, bottom), fill="#94a3b8", width=2)
    for i, (name, d1, d2, danger1, danger2) in enumerate(data):
        group_x = left + 165 + i * 390
        h1 = int(d1 / max_dist * 380)
        h2 = int(d2 / max_dist * 380)
        draw.rectangle((group_x, bottom - h1, group_x + 70, bottom), fill="#2563eb")
        draw.rectangle((group_x + 100, bottom - h2, group_x + 170, bottom), fill="#16a34a")
        draw.text((group_x - 18, bottom + 26), name, font=body_font, fill="#0f172a")
        draw.text((group_x - 5, bottom - h1 - 34), f"{d1:g}km", font=small_font, fill="#2563eb")
        draw.text((group_x + 82, bottom - h2 - 34), f"{d2:g}km", font=small_font, fill="#16a34a")
        draw.text((group_x - 22, bottom + 68), f"危险边 {danger1}", font=small_font, fill="#2563eb")
        draw.text((group_x + 92, bottom + 68), f"危险边 {danger2}", font=small_font, fill="#16a34a")
    draw.rectangle((1050, 160, 1085, 190), fill="#2563eb")
    draw.text((1100, 158), "普通最短路径", font=small_font, fill="#0f172a")
    draw.rectangle((1050, 205, 1085, 235), fill="#16a34a")
    draw.text((1100, 203), "安全路径", font=small_font, fill="#0f172a")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, quality=95)


def create_platform_arch_image(path: Path) -> None:
    w, h = 1600, 820
    img = Image.new("RGB", (w, h), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = load_image_font(42, True)
    body_font = load_image_font(25)
    small_font = load_image_font(21)
    draw.text((70, 42), "全国灾害响应平台架构示意图", font=title_font, fill="#0f172a")
    layers = [
        ("数据层", ["OSM 道路节点与边", "历史灾害事件", "交通拥堵映射"], "#dbeafe"),
        ("算法层", ["灾害影响叠加", "安全权重计算", "Dijkstra 路径规划"], "#dcfce7"),
        ("服务层", ["/api/scenarios", "/api/run-pipeline", "/artifact 图片读取"], "#fef3c7"),
        ("展示层", ["场景切换", "路线图展示", "流水线演示"], "#fee2e2"),
    ]
    y = 170
    for i, (name, items, color) in enumerate(layers):
        x = 100 + i * 365
        rounded_box(draw, (x, y, x + 285, y + 430), color, "#334155", 3)
        draw.text((x + 92, y + 30), name, font=body_font, fill="#0f172a")
        for j, item in enumerate(items):
            rounded_box(draw, (x + 28, y + 105 + j * 88, x + 257, y + 165 + j * 88), "#ffffff", "#cbd5e1", 2, 10)
            draw_wrapped_text(draw, (x + 45, y + 119 + j * 88), item, small_font, "#334155", 190)
        if i < len(layers) - 1:
            arrow(draw, (x + 295, y + 215), (x + 352, y + 215), "#334155", 4)
    draw.text((105, 690), "平台作用：把算法脚本输出转化为课堂可演示、可切换、可复现的交互界面。", font=small_font, fill="#475569")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, quality=95)


def create_ui_layout_image(path: Path) -> None:
    w, h = 1500, 840
    img = Image.new("RGB", (w, h), "#ffffff")
    draw = ImageDraw.Draw(img)
    title_font = load_image_font(42, True)
    body_font = load_image_font(24)
    small_font = load_image_font(21)
    draw.text((70, 42), "可视化页面三栏布局示意图", font=title_font, fill="#0f172a")
    rounded_box(draw, (80, 150, 430, 700), "#dbeafe", "#2563eb", 3)
    rounded_box(draw, (470, 150, 1030, 700), "#dcfce7", "#16a34a", 3)
    rounded_box(draw, (1070, 150, 1420, 700), "#fef3c7", "#d97706", 3)
    draw.text((185, 185), "左侧场景区", font=body_font, fill="#1d4ed8")
    draw.text((650, 185), "中间展示区", font=body_font, fill="#15803d")
    draw.text((1180, 185), "右侧指标区", font=body_font, fill="#b45309")
    left_items = ["全国灾害示意地图", "四川地震 / 北京洪水 / 上海火灾", "点击切换当前场景"]
    center_items = ["场景标题和摘要", "路线图或阶段图", "流水线步骤与运行日志"]
    right_items = ["节点和道路边数量", "普通路径与安全路径", "图例和风险说明"]
    for x, y, items in [(120, 260, left_items), (535, 260, center_items), (1115, 260, right_items)]:
        for i, item in enumerate(items):
            rounded_box(draw, (x, y + i * 105, x + 270 if x != 535 else x + 420, y + 70 + i * 105), "#ffffff", "#cbd5e1", 2, 10)
            draw_wrapped_text(draw, (x + 20, y + 18 + i * 105), item, small_font, "#334155", 230 if x != 535 else 370)
    draw.text((100, 745), "设计目标：让观众按照“选场景 - 看路线 - 读指标 - 理解流程”的顺序快速看懂项目。", font=small_font, fill="#475569")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, quality=95)


def create_pipeline_timeline_image(path: Path) -> None:
    create_flow_image(
        path,
        "快速演示模式七步流水线",
        ["选择灾害事件", "加载真实路网", "叠加灾害影响区", "标记风险拥堵", "普通最短路径", "安全救援路径"],
        accent="#16a34a",
    )


def generate_report_images() -> dict[str, Path]:
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "algorithm_flow": IMAGE_DIR / "algorithm_flow.png",
        "weight_formula": IMAGE_DIR / "weight_formula.png",
        "result_chart": IMAGE_DIR / "result_chart.png",
        "platform_arch": IMAGE_DIR / "platform_architecture.png",
        "ui_layout": IMAGE_DIR / "ui_layout.png",
        "pipeline_timeline": IMAGE_DIR / "pipeline_timeline.png",
    }
    create_flow_image(
        paths["algorithm_flow"],
        "数据与算法模块总体流程图",
        ["场景配置", "真实道路数据", "灾害影响映射", "安全权重计算", "Dijkstra 搜索", "结果输出"],
        accent="#2563eb",
    )
    create_weight_image(paths["weight_formula"])
    create_result_chart(paths["result_chart"])
    create_platform_arch_image(paths["platform_arch"])
    create_ui_layout_image(paths["ui_layout"])
    create_pipeline_timeline_image(paths["pipeline_timeline"])
    return paths


def set_run_font(run, size: float = 12, bold: bool = False, font: str = "宋体") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_font(paragraph, size: float = 12, bold: bool = False, font: str = "宋体") -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, font=font)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.5


def add_center_paragraph(doc: Document, text: str, size: float, bold: bool = False, spacing_after: float = 0) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_cover(doc: Document, name: str, student_id: str, role_title: str) -> None:
    add_center_paragraph(doc, "应急管理大学", 26)
    add_center_paragraph(doc, "结课报告书", 42, bold=True, spacing_after=80)
    for _ in range(4):
        doc.add_paragraph()

    lines = [
        "专    业                  资源与环境                   ",
        "学    院             计算机与信息安全学院              ",
        f"报告题目  {PROJECT_TITLE}   ",
        f"姓    名      {name}      学    号    {student_id}        ",
        "年    级      2025级     指导教师        袁静        ",
    ]
    for text in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=14, bold=True)
    doc.add_paragraph()
    add_center_paragraph(doc, REPORT_DATE, 14, bold=True)
    doc.add_page_break()


def add_static_toc(doc: Document, headings: list[tuple[int, str]]) -> None:
    add_center_paragraph(doc, "目录", 10.5)
    for level, title in headings:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(21 if level == 2 else 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        set_run_font(r, size=14 if level == 1 else 10.5, bold=(level == 1))
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_paragraph_font(p, size=16 if level == 1 else 14, bold=True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, size=12)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10)


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.text = ""
    r = p.add_run(text)
    set_run_font(r, size=12, bold=bold)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{name}"))
        if border is None:
            border = OxmlElement(f"w:{name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def add_picture_if_exists(doc: Document, path: Path, caption: str, width_cm: float = 14.5) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def report_one_sections() -> list[tuple[str, list[str]]]:
    return [
        ("1 摘要", [
            "本报告围绕“避开危险路段的救援路径规划”课程项目展开，重点记录本人在道路数据组织、灾害影响映射、图模型构建、Dijkstra 算法计算与结果分析等方面承担的工作。项目以真实道路网络和历史灾害事件为基础，将道路抽象为节点和边，将地震、洪水、火灾等灾害影响转化为道路风险属性，再使用 Dijkstra 算法分别计算普通最短路径和综合安全路径。",
            "本人主要负责数据与算法侧工作，包括梳理 OSM/高德道路数据结构、设计 nodes.csv、edges.csv、disaster_events.csv、road_disaster_mapping.csv 等核心数据表，完成道路边权重计算逻辑解释，整理四川汶川地震、北京暴雨洪涝、上海胶州路火灾三个场景的路径结果，并从危险边数量、综合代价和救援适用性角度进行分析。项目最终形成可运行的数据包、算法源码、路径对比结果和可视化结果，为课程汇报中的技术原理说明提供支撑。",
        ]),
        ("2 概述", [
            "灾害发生后，道路网络的通行条件会受到塌方、积水、火场管制和拥堵等因素影响。传统最短路径通常只关注距离或时间，而应急救援更关注能否安全、稳定、可靠地到达受灾点。因此，本项目不是简单做一张导航图，而是尝试把灾害风险引入路径规划模型。",
            "项目选取四川汶川地震、北京暴雨洪涝和上海胶州路火灾三个真实历史事件作为场景。三个场景分别对应山区地震、城市内涝和市中心消防救援，能够体现不同灾害条件下道路风险对路径选择的影响。",
        ]),
        ("3 需求分析", [
            "从本人负责的工作角度看，系统首先需要有真实道路数据，不能只依靠手工绘制的示意图。其次，灾害影响必须能够映射到道路边上，否则算法无法判断哪些路段更危险。再次，路径规划必须保持可解释性，老师能够清楚理解权重如何计算，为什么安全路径可能比最短路径更长。",
            "因此，本模块的核心需求可以归纳为四点：一是构建结构清晰的数据包；二是将灾害影响区与道路边做空间叠加；三是把距离、灾害风险和拥堵状态转化为非负边权重；四是使用 Dijkstra 算法输出可复核的路径结果。",
        ]),
        ("4 总体技术路线", [
            "本人参与梳理的数据与算法流程为：首先读取场景配置，确定救援起点、受灾终点、灾害事件位置和影响范围；然后读取或生成道路节点表和道路边表；接着根据灾害影响半径判断道路边是否进入危险区域；再叠加交通拥堵状态；最后根据不同权重模式运行 Dijkstra 算法。",
            "普通最短路径使用道路距离作为权重，安全路径使用 distance × risk_factor × (1 + congestion_weight × congestion) + fixed_cost 作为综合权重。这样既保持了 Dijkstra 算法对非负权重的要求，又能够让算法在风险较高的道路上产生更高代价，从而主动选择更安全的绕行路线。",
        ]),
        ("5 数据采集、整理与数据包构建", [
            "项目道路数据主要来自 OpenStreetMap / Overpass API 区域路网抽取，同时结合高德交通态势数据表达道路拥堵情况。道路节点保存为 nodes.csv，道路边保存为 edges.csv，历史灾害事件保存为 disaster_events.csv，道路与灾害叠加关系保存为 road_disaster_mapping.csv，交通状态映射保存为 road_traffic_mapping.csv。",
            "本人重点关注数据表之间的对应关系。nodes.csv 中的 node_id 是图搜索的基础，edges.csv 中的 from、to、distance、danger_type、congestion 是权重计算的直接输入，scenario.json 中保存风险系数和起终点信息，path_results.json 和 path_comparison.csv 则用于记录 Dijkstra 输出结果。",
        ]),
        ("6 图模型与权重设计", [
            "项目将道路网络抽象为无向加权图 G=(V,E)。节点 V 表示道路交叉点、道路折点、救援起点和受灾终点；边 E 表示相邻节点之间可通行的道路片段。每条边不仅保存距离，还保存道路名称、危险类型、拥堵程度、是否可通行和综合安全权重。",
            "本人负责将权重设计整理成可解释的工程逻辑。普通最短路径的权重为道路距离，安全路径的权重综合考虑距离、灾害风险系数、拥堵修正和固定惩罚。四川地震场景对 collapse 设置较高风险系数，北京洪水场景重点提高 flood 路段代价，上海火灾场景则对 fire 和 congestion 给予更强惩罚。",
        ]),
        ("7 Dijkstra 算法实现与结果输出", [
            "项目只使用 Dijkstra 算法，不引入 A* 或其他算法。算法实现中使用优先队列保存当前可扩展节点，每次取出总代价最小的节点进行松弛更新，直到确定目标节点。由于本项目所有边权重均为非负数，因此 Dijkstra 能够保证在当前图模型和当前权重定义下得到全局最小代价路径。",
            "本人整理了普通最短路径与安全路径两类结果。普通最短路径用于模拟传统距离优先路线，安全路径用于体现灾害救援场景下的综合安全选择。输出结果包括路径节点序列、总距离、综合代价、危险边数量、危险类型和救援适用性评价。",
        ]),
        ("8 三类场景结果分析", [
            "四川汶川地震场景中，普通最短路径会经过大量 collapse 和 congestion 路段，而安全路径距离明显增加，但危险边数量降为 0，说明权重设置能够有效引导路径绕开主要地震影响区。",
            "北京暴雨洪涝场景中，安全路径相比普通最短路径只增加少量距离，但显著减少了 flood 和 congestion 风险暴露，说明城市道路网络提供了较多替代路线。",
            "上海胶州路火灾场景中，由于目标点本身位于火灾影响区域，安全路径无法完全避开 fire 风险，但能够减少进入核心风险区的路段数量，更符合消防救援必须接近火点但应尽量降低风险暴露的逻辑。",
        ]),
        ("9 全套交付成果清单", [
            "本人负责整理的数据与算法侧成果主要包括：三类场景的数据表、Dijkstra 路径结果、路径对比表、边权重表、灾害映射说明、权重公式说明和结果分析文字。相关文件集中在 data、outputs、src 和 docs 目录中。",
            "其中 src/rescue_planner.py 是图构建、权重计算和 Dijkstra 路径规划的核心程序；outputs/osm_* 目录保存路径结果和可视化图；docs/项目14_避开危险路段的救援路径规划_汇报文档.md 用于支撑汇报说明。",
        ]),
        ("10 实验总结与收获", [
            "通过本项目，本人更加清楚地理解了图模型、边权重和最短路径算法之间的关系。Dijkstra 算法本身并不复杂，真正影响路径结果的是如何把真实道路、灾害风险和拥堵状态转化为合理的边权重。",
            "本次工作最大的收获是把抽象算法放入灾害救援业务场景中解释。安全路径不是单纯追求更短，而是在当前风险模型下让综合代价最低。后续如果有更精确的道路损毁、积水深度和实时封控数据，可以进一步校准风险系数，使模型更接近真实应急场景。",
        ]),
        ("11 附录", [
            "附录主要列出本人负责模块相关的关键公式、数据字段和结果图，便于答辩时快速说明算法依据和数据来源。",
        ]),
    ]


def report_two_sections() -> list[tuple[str, list[str]]]:
    return [
        ("1 摘要", [
            "本报告围绕“避开危险路段的救援路径规划”课程项目展开，重点记录本人在可视化页面、演示平台、流水线展示、报告材料整理和最终成果集成方面承担的工作。项目在完成道路数据、灾害映射和 Dijkstra 算法计算的基础上，进一步实现了全国灾害响应平台式的网页演示界面，使老师能够直观看到不同灾害场景下的路径规划过程。",
            "本人主要负责把后端生成的数据结果转化为可展示、可解释、适合课堂汇报的界面和图像。具体工作包括梳理三类灾害场景信息，设计页面布局和交互流程，接入 route_map_abstract.png 和 pipeline_step_1.png 至 pipeline_step_7.png 等可视化产物，配置快速演示与真实运行两种模式，并整理最终汇报文档、PPT 说明和演示口径。通过该部分工作，项目从单纯算法脚本扩展为可以完整展示“数据获取 - 灾害叠加 - 路径计算 - 可视化输出”的原型系统。",
        ]),
        ("2 概述", [
            "救援路径规划项目如果只有 CSV 表和算法代码，虽然可以证明计算过程存在，但课堂展示时不够直观。为了让项目更容易理解，需要把抽象的图模型、危险边、普通最短路径和安全路径通过界面呈现出来。",
            "本人负责的重点是平台化展示和结果集成。页面左侧展示全国灾害场景分布，中间展示路线图和流水线阶段图，右侧展示数据规模、路径结果和图例说明。这样的组织方式能够让观众快速理解项目做了什么、数据从哪里来、算法如何运行、最终路径为什么不同。",
        ]),
        ("3 需求分析", [
            "可视化部分首先需要清晰表达三类灾害场景。四川地震、北京洪水和上海火灾在空间范围、风险类型和路径逻辑上都不同，因此页面需要支持场景切换，并在每个场景下展示对应的起点、终点、历史灾害事件、道路风险和 Dijkstra 结果。",
            "其次，平台需要区分快速演示和真实运行。快速演示适合课堂录屏，直接播放已生成的阶段图；真实运行则调用本地脚本重新执行 Dijkstra 和可视化渲染，用于证明系统不是静态图片拼接，而是具有可复现的计算流程。",
        ]),
        ("4 总体技术路线", [
            "本人参与设计的可视化技术路线为：后端读取 data/osm_* 与 outputs/osm_* 目录中的结果文件，整理成 JSON 接口；前端通过 /api/scenarios 获取场景数据，再根据用户选择切换地图、指标、路径图和流水线步骤；当用户点击真实运行按钮时，后端调用 rescue_planner.py 和 create_abstract_route_maps.py 重新生成结果。",
            "为了保证展示效果，项目将可视化分为两类。一类是最终路径图 route_map_abstract.png，用于展示普通最短路径与安全路径对比；另一类是 pipeline_step_1.png 到 pipeline_step_7.png，用于展示从选择灾害事件到输出最终结果的完整流程。",
        ]),
        ("5 页面结构与交互设计", [
            "平台页面采用三栏结构。左侧为场景选择区，包含全国示意地图和灾害场景列表；中间为主要展示区，包含场景标题、结果路线图、流水线步骤和日志窗口；右侧为指标说明区，展示节点数量、道路边数量、风险映射数量、交通映射数量、普通最短路径和安全路径结果。",
            "这种结构的优点是信息层次清楚。观众先在左侧选择场景，再在中间看到路径图，最后在右侧读取数据和结果指标。页面不需要大量说明文字，主要通过图像、指标和步骤状态引导观看。",
        ]),
        ("6 流水线演示功能实现", [
            "快速演示功能按照七个阶段播放：选择历史灾害事件和救援起终点、加载真实道路网络、叠加历史灾害影响区、标记危险和拥堵路段、运行 Dijkstra 得到普通最短路径、运行 Dijkstra 得到安全救援路径、输出路径对比表和可视化路线图。",
            "本人在整理该部分时重点保证步骤和报告中的技术路线一致。每个步骤不仅是界面动画，也对应项目真实处理流程，使老师在观看演示时能够把页面变化和算法过程联系起来。",
        ]),
        ("7 可视化图像与图例说明", [
            "最终路线图采用颜色分层方式表达信息。蓝色线表示普通最短路径，绿色线表示安全路径，红色或橙色表示灾害或拥堵风险路段，浅灰色道路表示周边真实道路背景。这样既保留真实道路形态，又避免所有道路同等显示造成视觉混乱。",
            "在图例说明中，本人重点区分了“背景道路”和“参与路径规划的道路”。部分道路只作为可视化背景，用于增强地图语境，并不直接参与 Dijkstra 搜索。这样可以避免为了画面丰富而破坏算法解释。",
        ]),
        ("8 三类场景展示效果", [
            "四川汶川地震场景中，页面重点展示山区地震影响区和塌方风险路段，安全路径绕开主要风险区域，能够直观体现“距离更长但风险更低”的救援逻辑。",
            "北京暴雨洪涝场景中，页面重点展示积水与拥堵对城市道路的影响，安全路径相比普通最短路径绕行幅度较小，但危险边数量明显减少，适合说明城市内涝下的路径优化意义。",
            "上海胶州路火灾场景中，页面重点展示火点周边风险和拥堵，安全路径无法完全远离目标点附近风险，但能减少穿越火场核心区的路段，体现消防救援场景的特殊性。",
        ]),
        ("9 全套交付成果清单", [
            "本人负责整理和集成的成果包括 national_disaster_platform/index.html、static/app.js、static/styles.css、server.py、三类场景可视化输出、流水线阶段图、汇报文档和演示说明。通过这些内容，项目可以从命令行脚本运行扩展到网页交互展示。",
            "最终交付材料包括源码、数据包、输出图像、路径结果表、网页平台、汇报文档和演示脚本。可视化平台是连接算法实现和课堂汇报的桥梁，使项目成果更容易被理解和验收。",
        ]),
        ("10 实验总结与收获", [
            "通过本项目，本人认识到可视化并不是简单美化页面，而是要把算法过程讲清楚。路径规划中有很多抽象概念，例如边权重、危险映射、综合代价和 Dijkstra 搜索，如果没有合适的界面和图像，观众很难快速理解。",
            "本次工作最大的收获是学会把技术结果组织成可汇报、可演示、可复现的形式。后续如果继续完善，可以增加地图底图联动、实时交通状态刷新、更多灾害类型切换和路径参数调节，使平台更接近真实应急指挥辅助系统。",
        ]),
        ("11 附录", [
            "附录主要列出本人负责模块相关的界面文件、接口路径、流水线步骤图和关键展示截图，便于答辩时说明平台功能和成果位置。",
        ]),
    ]


def make_report(
    name: str,
    student_id: str,
    role_title: str,
    sections: list[tuple[str, list[str]]],
    out_path: Path,
    images: dict[str, Path],
) -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc, name, student_id, role_title)
    toc_headings = []
    for heading, _ in sections:
        toc_headings.append((1, heading))
        if heading.startswith("1 "):
            toc_headings.append((2, "1.1成员分工说明"))
        if heading.startswith("2 "):
            toc_headings.extend([(2, "2.1 研究背景"), (2, "2.2 整体研究目标"), (2, "2.3 核心重难点"), (2, "2.4 建设原则")])
    add_static_toc(doc, toc_headings)

    for heading, paragraphs in sections:
        add_heading(doc, heading, 1)
        if heading == "1 摘要":
            for text in paragraphs:
                add_body(doc, text)
            add_heading(doc, "1.1成员分工说明", 2)
            add_body(doc, "小组成员：成员一、成员二")
            add_table(
                doc,
                "表1-1 成员分工内容",
                ["成员", "负责模块", "核心工作内容"],
                [
                    ["成员一", "数据、图模型与 Dijkstra 算法模块", "负责道路数据整理、灾害影响区映射、图结构构建、边权重设计、Dijkstra 普通最短路径与安全路径计算、结果表与算法章节整理。"],
                    ["成员二", "可视化平台、流水线演示与成果集成模块", "负责全国灾害响应平台页面、场景切换、路径图展示、快速演示、真实运行接口、汇报材料与最终展示成果整理。"],
                ],
            )
            continue

        if heading == "2 概述":
            add_heading(doc, "2.1 研究背景", 2)
            add_body(doc, paragraphs[0])
            add_heading(doc, "2.2 整体研究目标", 2)
            add_body(doc, paragraphs[1])
            add_heading(doc, "2.3 核心重难点", 2)
            add_body(doc, "项目难点主要体现在真实路网规模较大、灾害数据与道路数据形式不同、权重设计需要可解释、最终结果还要适合课堂汇报展示。本人负责模块围绕这些难点，将工程实现、数据结果和可视化表达统一起来。")
            add_heading(doc, "2.4 建设原则", 2)
            add_body(doc, "项目坚持使用真实道路或真实区域路网数据，坚持只使用 Dijkstra 算法，坚持高德 Key 不写入项目文件，坚持把危险路段解释为灾害影响区与道路空间叠加得到的结果。")
            continue

        for text in paragraphs:
            add_body(doc, text)

        if "算法" in role_title or "数据" in role_title:
            if heading.startswith("4 "):
                add_picture_if_exists(doc, images["algorithm_flow"], "图4-1 数据与算法模块总体流程图", width_cm=15.2)
            if heading.startswith("6 "):
                add_picture_if_exists(doc, images["weight_formula"], "图6-1 安全路径综合权重设计示意图", width_cm=14.8)
            if heading.startswith("8 "):
                add_picture_if_exists(doc, images["result_chart"], "图8-1 三类灾害场景路径结果对比图", width_cm=14.8)
        else:
            if heading.startswith("4 "):
                add_picture_if_exists(doc, images["platform_arch"], "图4-1 全国灾害响应平台架构示意图", width_cm=15.2)
            if heading.startswith("5 "):
                add_picture_if_exists(doc, images["ui_layout"], "图5-1 可视化页面三栏布局示意图", width_cm=14.8)
            if heading.startswith("6 "):
                add_picture_if_exists(doc, images["pipeline_timeline"], "图6-1 快速演示模式七步流水线", width_cm=15.2)

        if heading.startswith("8 "):
            add_table(
                doc,
                "表8-1 三类灾害场景路径结果对比",
                ["场景", "普通最短路径特点", "安全路径特点"],
                [
                    ["四川汶川地震", "距离较短但经过较多 collapse 风险边", "距离增加但可绕开主要塌方风险"],
                    ["北京暴雨洪涝", "存在 flood 与 congestion 风险暴露", "小幅绕行即可显著减少危险边"],
                    ["上海胶州路火灾", "接近火点，fire 风险较集中", "在保证到达火点前提下降低核心风险暴露"],
                ],
            )

        if heading.startswith("11 "):
            if "算法" in role_title or "数据" in role_title:
                add_picture_if_exists(doc, ROOT / "outputs" / "osm_sichuan_earthquake" / "route_map_abstract.png", "图11-1 四川汶川地震场景路径规划结果")
                add_body(doc, "关键公式：safe_weight = distance × risk_factor × (1 + congestion_weight × congestion) + fixed_cost。该公式用于把道路长度、灾害风险和拥堵状态统一为 Dijkstra 可处理的非负权重。")
            else:
                add_picture_if_exists(doc, ROOT / "outputs" / "osm_beijing_flood" / "pipeline_step_7.png", "图11-1 北京洪水场景流水线最终结果图")
                add_body(doc, "关键接口：/api/scenarios 用于返回三类场景数据，/api/run-pipeline/{sid} 用于触发真实运行模式，/artifact 用于读取本地可视化图片。")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    images = generate_report_images()
    make_report(
        "成员一（待填写）",
        "学号待填写",
        "数据与算法实现",
        report_one_sections(),
        OUT_DIR / "成员一-数据与算法实现-结课报告.docx",
        images,
    )
    make_report(
        "成员二（待填写）",
        "学号待填写",
        "可视化平台与成果集成",
        report_two_sections(),
        OUT_DIR / "成员二-可视化平台与成果集成-结课报告.docx",
        images,
    )
    print(OUT_DIR)


if __name__ == "__main__":
    main()
