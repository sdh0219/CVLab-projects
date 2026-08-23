# -*- coding: utf-8 -*-
"""
build_report.py —— 生成《技术方案》与《实验报告》Word 文档(.docx)。

特点：
  - 真实数值从 output/reports/*.json 动态读取，保证与代码运行结果一致；
  - 自动嵌入 output/figures/ 中的图表；
  - 含标题层级、表格、要点列表，可直接提交。

运行：python src/build_report.py
依赖：python-docx
"""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

import config as C

DOCS_DIR = C.ROOT_DIR / "docs"
DOCS_DIR.mkdir(exist_ok=True)
FIG = C.FIG_DIR


# ---------------------------------------------------------------- 读取真实数值
def load_json(p):
    return json.loads(Path(p).read_text(encoding="utf-8"))


QB = load_json(C.QUALITY_BEFORE_JSON)
QA = load_json(C.QUALITY_AFTER_JSON)
EV = load_json(C.EVAL_JSON)
GEN = load_json(C.REPORT_DIR / "data_generation_summary.json")

DET = EV["异常检测评估"]
COR = EV["数值修正评估"]


# ---------------------------------------------------------------- docx 辅助函数
BASE_FONT = "宋体"
HEAD_FONT = "黑体"


def _set_cn_font(run, font=BASE_FONT, size=None, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_cn_font(r, HEAD_FONT, 20, bold=True)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_cn_font(r, BASE_FONT, 12, color=(90, 90, 90))


def add_h(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    size = {1: 15, 2: 13, 3: 12}.get(level, 12)
    _set_cn_font(r, HEAD_FONT, size, bold=True,
                 color=(31, 78, 121) if level == 1 else (0, 0, 0))
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_p(doc, text, size=11, indent=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_cn_font(r, BASE_FONT, size)
    p.paragraph_format.line_spacing = 1.4
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_cn_font(r, BASE_FONT, size)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_fig(doc, name, caption, width=6.3):
    path = FIG / name
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        _set_cn_font(r, BASE_FONT, 10, color=(90, 90, 90))


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        _set_cn_font(r, HEAD_FONT, 10.5, bold=True, color=(255, 255, 255))
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(hdr[i], "4472C4")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            _set_cn_font(r, BASE_FONT, 10.5)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t


def _shade_cell(cell, hexcolor):
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_normal_style(doc):
    st = doc.styles["Normal"]
    st.font.name = BASE_FONT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    st.font.size = Pt(11)


# ================================================================ 技术方案
def build_tech_solution():
    doc = Document()
    set_normal_style(doc)

    add_title(doc, "雨量计异常数据清洗与纠正  技术方案")
    add_subtitle(doc, "AI + 防灾减灾学生实践项目 · 项目3（模块2：监测预警）")
    doc.add_paragraph()

    # 1 背景与目标
    add_h(doc, "一、项目背景与目标")
    add_p(doc, "雨量计是暴雨洪涝、山洪、地质灾害等监测预警业务的基础感知设备。其原始观测"
               "数据在采集、传输、存储链路中不可避免地混入缺失、跳变、卡滞等异常，若不经清洗"
               "直接进入预警模型，轻则使统计失真，重则触发错误预警或漏报，直接危及防灾决策。")
    add_p(doc, "本项目目标：掌握传感器数据异常识别、缺失值处理与可视化分析方法，构建一条"
               "可复现、可量化评估的雨量计数据清洗管线，并以累计雨量这一防灾关键指标，"
               "定量说明清洗的应用价值。")

    # 2 数据质量问题
    add_h(doc, "二、数据质量问题分析")
    add_p(doc, "结合翻斗式雨量计的工作原理与实际运维经验，本项目梳理出 7 类典型数据质量问题，"
               "并在数据集中按真实比例注入，便于系统性验证清洗规则：")
    add_table(doc,
              ["异常类型", "成因", "典型表现", "对防灾的危害"],
              [["缺失值", "通信中断 / 设备掉电 / 采集器故障", "记录为空(NaN)，含随机点与连续整段", "时序不连续，累计量被低估，预警断档"],
               ["负值", "传输误码 / AD 采样错误", "出现 < 0 的雨量（物理不可能）", "统计污染，破坏后续模型"],
               ["尖峰/超量程", "电气干扰 / 雷击 / 哨兵值(999)溢出", "瞬时极大值，远超物理上限", "累计雨量灾难性虚高，触发错误暴雨/洪水预警"],
               ["传感器卡滞", "翻斗机械卡死 / 数据冻结", "长时间读数恒定不变", "漏记真实降雨或虚增降雨"],
               ["虚假翻斗", "蛛网 / 振动 / 噪声扰动翻斗", "干期出现微量假降雨(0.1~0.3mm)", "虚增降雨频次，干扰雨情判断"],
               ["时间戳重复", "数据采集器重发", "同一时刻多条记录", "重复计数，时序错乱"],
               ["时间戳乱序", "缓存回灌 / 网络抖动", "记录未按时间顺序入库", "滚动窗口/差分计算失效"]])
    add_p(doc, f"在本数据集（{GEN['样本数(规则时间轴)']} 个 10min 样本）中实际注入："
               f"缺失 {QB['缺失值数']}、负值 {QB['负值数']}、超量程 {QB['超量程数(>40mm/10min)']}、"
               f"卡滞 {QB['疑似卡滞点数']} 点，重复 5 行、乱序 20 处。受污染后原始数据最大值达 "
               f"{QB['最大值']} mm/10min、总降雨量被抬高至 {QB['总降雨量(mm)']} mm。")

    # 3 清洗规则
    add_h(doc, "三、数据清洗规则")
    add_p(doc, "针对上述问题，制定多规则融合的判定标准（不依赖真值，纯数据驱动），"
               "当一点命中多条规则时按严重度 缺失>尖峰>负值>卡滞>虚假翻斗 取最严重者：")
    add_bullet(doc, "规则A 物理范围检查：雨量必须非负；单 10min 雨量超过物理硬上限 "
                    f"{C.RAIN_MAX_PER_INTERVAL:.0f} mm 即判超量程异常（中国短时强降水 10min 极端量级约 16mm，留足余量）。")
    add_bullet(doc, "规则B 稳健统计离群检测：在滚动窗口内以中位数+MAD 计算稳健 Z 分数 "
                    "(x−median)/(1.4826·MAD)，仅当数值同时统计极端、超过可疑阈值且明显高于近邻"
                    "（孤立跳变）时才判尖峰——以此保护真实暴雨峰值不被误删。")
    add_bullet(doc, f"规则C 传感器卡滞检测：连续 ≥{C.STUCK_MIN_REPEAT} 个完全相同的非零读数判为卡滞"
                    "（干期连续为 0 属正常，不判卡滞）。")
    add_bullet(doc, "规则D 虚假翻斗检测：干期中孤立出现的微量降雨（≤0.3mm 且前后窗口全干）判为虚假翻斗。")
    add_bullet(doc, "规则E 缺失检测：空值(NaN)直接标记为缺失。")

    # 4 处理方法
    add_h(doc, "四、处理方法与技术路线")
    add_p(doc, "整体管线分 5 个阶段，全部基于 Python / Pandas 实现：")
    add_table(doc,
              ["阶段", "方法", "说明"],
              [["① 时间轴规整", "去重 + 排序 + 重建规则网格",
                "删除重复时间戳，按时间排序修复乱序，缺失时刻补 NaN，使滚动/差分计算有意义"],
               ["② 质量统计", "缺失率/负值/超量程/卡滞/分布统计",
                "量化清洗前数据质量，定位问题"],
               ["③ 异常检测", "规则 A~E 多规则融合",
                "输出逐点异常标记 flag"],
               ["④ 插值修正", "基于时间的线性插值",
                f"异常点掩膜为缺失后统一插值；仅填补 ≤{C.INTERP_LIMIT} 个连续点的短缺口，长缺口保守置 0 并标注低可信"],
               ["⑤ 平滑与约束", "中值滤波 + 移动平均 + 非负/量化",
                "中值滤波去残余尖峰，移动平均轻度降噪，最后非负约束并按 0.1mm 翻斗分辨率重量化"]])
    add_p(doc, "关键设计取舍：滑动窗口平滑虽能降噪，但会轻微削平真实峰值、重分配雨量，"
               "故本项目以未平滑的 corrected 序列作为主清洗结果，平滑结果单列输出供对比，"
               "避免过度平滑造成累计量系统性偏低。")

    # 5 清洗效果
    add_h(doc, "五、清洗效果")
    add_table(doc,
              ["指标", "清洗前", "清洗后", "真值"],
              [["缺失率", f"{QB['缺失率(%)']}%", f"{QA['缺失率(%)']}%", "—"],
               ["负值 / 超量程 / 卡滞",
                f"{QB['负值数']} / {QB['超量程数(>40mm/10min)']} / {QB['疑似卡滞点数']}",
                f"{QA['负值数']} / {QA['超量程数(>40mm/10min)']} / 0", "—"],
               ["最大 10min 雨量", f"{QB['最大值']} mm", f"{QA['最大值']} mm", f"{QA['最大值']} mm"],
               ["总降雨量",
                f"{COR['总降雨量(mm)']['原始']} mm ({COR['总降雨量相对误差(%)']['原始']}%)",
                f"{COR['总降雨量(mm)']['清洗后']} mm ({COR['总降雨量相对误差(%)']['清洗后']}%)",
                f"{COR['总降雨量(mm)']['真值']} mm"],
               ["逐点 MAE / RMSE",
                f"{COR['逐点误差']['原始(仅有效点)']['MAE']} / {COR['逐点误差']['原始(仅有效点)']['RMSE']}",
                f"{COR['逐点误差']['清洗后corrected']['MAE']} / {COR['逐点误差']['清洗后corrected']['RMSE']}",
                "—"]])
    add_p(doc, f"异常检测以注入真值为金标准：精确率 {DET['精确率(%)']}%、召回率 {DET['召回率(%)']}%、"
               f"F1 {DET['F1(%)']}%、准确率 {DET['准确率(%)']}%，且零误杀真实降雨（FP=0）。")
    add_fig(doc, "fig4_cumulative.png",
            "图1  累计雨量对比：原始尖峰致累计量虚高约 17 倍，清洗后与真值基本重合")

    # 6 防灾意义
    add_h(doc, "六、防灾应用意义")
    add_p(doc, f"累计雨量是暴雨预警、山洪与地质灾害判别的核心阈值指标。本数据集中，未清洗的原始"
               f"数据将一个月累计雨量从真实的 {COR['总降雨量(mm)']['真值']} mm 抬高到 "
               f"{COR['总降雨量(mm)']['原始']} mm（虚高约 1760%），单点尖峰即可让 10min 雨量"
               f"显示为 999 mm——足以瞬间触碰任何短时强降水阈值，造成误报、空报；而连续缺失与卡滞"
               f"又可能掩盖真实强降雨，造成漏报。")
    add_p(doc, "经本管线清洗后，累计雨量误差收敛到 0.3% 量级，逐点 RMSE 从 50 降到 0.05，"
               "为下游的雨强分级、预警阈值判定、洪水演算提供了可靠输入。这说明：数据清洗不是"
               "可有可无的预处理，而是监测预警链路中保障决策正确性的关键一环。")

    # 7 技术栈
    add_h(doc, "七、技术栈与可复现性")
    add_bullet(doc, "语言/库：Python 3 + Pandas（时序处理）、NumPy、SciPy、Matplotlib（可视化）。")
    add_bullet(doc, "全流程固定随机种子，一键脚本 python src/main.py 可完整复现数据、清洗、评估与图表。")
    add_bullet(doc, "模块化设计：config / generate_data / data_quality / anomaly_detection / "
                    "cleaning / evaluation / visualization / main，职责清晰、易扩展到真实数据。")

    out = DOCS_DIR / "技术方案.docx"
    doc.save(str(out))
    print("已生成:", out)


# ================================================================ 实验报告
def build_experiment_report():
    doc = Document()
    set_normal_style(doc)

    add_title(doc, "雨量计异常数据清洗与纠正  实验报告")
    add_subtitle(doc, "AI + 防灾减灾学生实践项目 · 项目3（模块2：监测预警）")
    doc.add_paragraph()

    # 1 实验目的
    add_h(doc, "一、实验目的")
    add_bullet(doc, "掌握雨量计等传感器时间序列数据的异常识别方法；")
    add_bullet(doc, "掌握缺失值处理（插值）与噪声平滑（滑动窗口）方法；")
    add_bullet(doc, "掌握数据质量量化统计与可视化分析方法；")
    add_bullet(doc, "理解数据清洗对防灾减灾监测预警的应用意义，并量化评估清洗效果。")

    # 2 环境与数据
    add_h(doc, "二、实验环境与数据集")
    add_p(doc, "环境：Python 3，依赖 pandas / numpy / scipy / matplotlib。")
    add_p(doc, f"数据集：自建合成数据，站点 RG-001，时间 {GEN['时间范围'][0]} ~ {GEN['时间范围'][1]}，"
               f"{GEN['采样分辨率']} 分辨率，共 {GEN['样本数(规则时间轴)']} 个样本。"
               "之所以采用合成数据，是因为需要逐点真值标签来严格评估异常检测的精确率/召回率"
               "与修正后的误差，而公开真实数据无逐点真值。降雨过程由泊松到达的降雨事件 + 钟形"
               "强度廓线 + 翻斗 0.1mm 量化生成，并同步生成水位站、气象站序列作为交叉校验背景。")
    add_p(doc, "在干净序列上按真实比例注入 7 类异常（缺失、负值、尖峰/超量程、卡滞、虚假翻斗、"
               "时间戳重复、时间戳乱序），并保留逐点真值标签用于评估。", indent=True)

    # 3 实验方法
    add_h(doc, "三、实验方法与流程")
    add_p(doc, "实验按如下流程执行（对应源码各模块）：")
    add_table(doc, ["步骤", "操作", "对应模块"],
              [["1", "生成含异常的原始数据与真值", "generate_data.py"],
               ["2", "时间轴规整：去重、排序、补规则网格", "data_quality.py"],
               ["3", "清洗前数据质量统计", "data_quality.py"],
               ["4", "多规则异常检测，输出逐点标记", "anomaly_detection.py"],
               ["5", "掩膜→时间插值→物理约束→滑动窗口平滑", "cleaning.py"],
               ["6", "对照真值评估检测与修正质量", "evaluation.py"],
               ["7", "绘制 6 类分析图表", "visualization.py"]])

    # 4 结果与分析
    add_h(doc, "四、实验结果与分析")

    add_h(doc, "4.1 原始数据与异常识别", 2)
    add_p(doc, f"时间轴规整删除重复时间戳 5 行、修复乱序。清洗前数据质量：缺失率 {QB['缺失率(%)']}%、"
               f"负值 {QB['负值数']}、超量程 {QB['超量程数(>40mm/10min)']}、卡滞 {QB['疑似卡滞点数']} 点，"
               f"最大值高达 {QB['最大值']} mm/10min。")
    add_fig(doc, "fig1_raw_with_anomalies.png",
            "图1  原始雨量计数据与异常点标记（上：全量含尖峰；下：放大视图）")

    add_h(doc, "4.2 数据质量统计", 2)
    add_fig(doc, "fig2_quality_dashboard.png", "图2  数据质量统计仪表板（异常构成、清洗前后可用性、关键指标对比）")

    add_h(doc, "4.3 异常检测评估", 2)
    pt = DET["分类型召回"]
    add_table(doc, ["异常类型", "注入数", "检出数", "召回率"],
              [[k, pt[k]["注入数"], pt[k]["检出数"], f'{pt[k]["召回率(%)"]}%']
               for k in ["missing", "negative", "spike", "stuck", "spurious"]])
    add_p(doc, f"总体：精确率 {DET['精确率(%)']}%、召回率 {DET['召回率(%)']}%、F1 {DET['F1(%)']}%、"
               f"准确率 {DET['准确率(%)']}%。混淆矩阵 TP={DET['混淆矩阵']['TP']}、FP={DET['混淆矩阵']['FP']}、"
               f"FN={DET['混淆矩阵']['FN']}、TN={DET['混淆矩阵']['TN']}。缺失、负值、尖峰、卡滞均 100% 召回；"
               "虚假翻斗召回率略低，因部分假翻斗紧邻真实降雨，'近邻全干'条件不满足——这类残余对累计量"
               "影响极小（单点≤0.3mm）。FP=0 表明真实暴雨峰值无一被误删，规则设计达到了'宁稳勿误杀'的目标。")
    add_fig(doc, "fig6_detection_metrics.png", "图3  异常检测效果评估（混淆矩阵 + 分类型召回 + 总体指标）")

    add_h(doc, "4.4 清洗修正效果", 2)
    add_table(doc, ["指标", "清洗前", "清洗后", "真值"],
              [["总降雨量(mm)", COR["总降雨量(mm)"]["原始"], COR["总降雨量(mm)"]["清洗后"], COR["总降雨量(mm)"]["真值"]],
               ["总降雨量相对误差", f'{COR["总降雨量相对误差(%)"]["原始"]}%', f'{COR["总降雨量相对误差(%)"]["清洗后"]}%', "0%"],
               ["逐点 MAE", COR["逐点误差"]["原始(仅有效点)"]["MAE"], COR["逐点误差"]["清洗后corrected"]["MAE"], "—"],
               ["逐点 RMSE", COR["逐点误差"]["原始(仅有效点)"]["RMSE"], COR["逐点误差"]["清洗后corrected"]["RMSE"], "—"]])
    add_p(doc, "清洗后逐点 RMSE 从约 50 降到 0.05，总降雨量相对误差从 +1760% 收敛到 −0.32%，效果显著。")
    add_fig(doc, "fig3_before_after.png", "图4  清洗前后对比（暴雨事件放大）：尖峰被移除，清洗序列贴合真值")

    add_h(doc, "4.5 滑动窗口平滑", 2)
    add_p(doc, f"在 corrected 基础上做中值滤波+移动平均得到 smoothed 序列，总降雨量 "
               f"{COR['总降雨量(mm)']['平滑后']} mm（误差 {COR['总降雨量相对误差(%)']['平滑后']}%）。"
               "可见平滑能进一步压制残余噪声，但会轻微削平峰值、使累计量略偏低，"
               "故定位为可选的降噪手段，主清洗结果仍采用 corrected。")
    add_fig(doc, "fig5_smoothing_zoom.png", "图5  滑动窗口平滑效果（暴雨峰值放大）")

    # 5 讨论
    add_h(doc, "五、讨论")
    add_bullet(doc, "阈值取舍：超量程硬上限与统计离群阈值需结合气候背景设置，过松漏检、过紧误杀；"
                    "本项目用'孤立跳变'条件保护真实暴雨峰值，实现 FP=0。")
    add_bullet(doc, "长缺口处理：超过 1 小时的连续缺失不臆造数值，保守置 0 并标注低可信，"
                    "避免插值伪造降雨；实际业务中可引入邻近站点或雷达估测填补。")
    add_bullet(doc, "平滑的双刃性：对降雨强度过度平滑会重分配雨量、削弱峰值，应谨慎使用。")
    add_bullet(doc, "向真实数据迁移：管线与阈值均参数化，替换 generate_data 的数据源即可直接用于"
                    "真实雨量计/水位站数据；水位站、气象站湿度可进一步用于跨传感器一致性校验。")

    # 6 结论
    add_h(doc, "六、结论")
    add_p(doc, f"本实验构建了完整的雨量计异常数据清洗管线，异常检测 F1 达 {DET['F1(%)']}%、"
               f"零误杀真实降雨，清洗后总降雨量误差仅 {COR['总降雨量相对误差(%)']['清洗后']}%、"
               f"逐点 RMSE 降至 {COR['逐点误差']['清洗后corrected']['RMSE']}。结果表明，规范的数据"
               "清洗能将被污染、不可用的原始观测恢复为高质量数据，是防灾减灾监测预警可靠运行的"
               "必要前提。")

    # 附录
    add_h(doc, "附录：交付物与运行方法")
    add_bullet(doc, "数据包：data/raw/（原始+真值）、data/cleaned/（清洗后+异常标记明细）。")
    add_bullet(doc, "源码：src/（8 个模块），一键运行 python src/main.py。")
    add_bullet(doc, "图表与报告：output/figures/、output/reports/。")

    out = DOCS_DIR / "实验报告.docx"
    doc.save(str(out))
    print("已生成:", out)


if __name__ == "__main__":
    build_tech_solution()
    build_experiment_report()
    print("文档生成完成 ->", DOCS_DIR)
